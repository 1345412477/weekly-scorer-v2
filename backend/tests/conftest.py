"""测试配置和通用 fixtures"""
import os
import sys
import tempfile
from datetime import date, timedelta

import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.database import Base, get_db
from app.main import app
from app.models.models import Department, Person, ScoringConfig
from app.core.auth import ensure_default_admin

TEST_DB_URL = "sqlite+aiosqlite:///./test_weekly_scorer.db"
test_engine = create_async_engine(TEST_DB_URL, echo=False)
TestSessionLocal = async_sessionmaker(test_engine, class_=AsyncSession, expire_on_commit=False)


async def override_get_db():
    async with TestSessionLocal() as session:
        yield session


app.dependency_overrides[get_db] = override_get_db


@pytest_asyncio.fixture(autouse=True)
async def setup_database():
    # 清除 AI 评分器的全局缓存，避免跨测试用例残留
    from app.services.ai_scorer import _clear_db_model_cache
    _clear_db_model_cache()

    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest_asyncio.fixture
async def db():
    async with TestSessionLocal() as session:
        yield session


@pytest_asyncio.fixture
async def client():
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c


@pytest_asyncio.fixture
async def admin_headers(client, db):
    await ensure_default_admin(db)
    resp = await client.post(
        "/api/v1/auth/login",
        json={"username": "admin", "password": "admin123"},
    )
    assert resp.status_code == 200
    return {"Authorization": f"Bearer {resp.json()['access_token']}"}


@pytest_asyncio.fixture
async def seed_departments(db):
    depts = [
        Department(id="dept-tech", name="技术部", description="技术研发"),
        Department(id="dept-product", name="产品部", description="产品设计"),
        Department(id="dept-sales", name="销售部", description="市场营销"),
    ]
    for d in depts:
        db.add(d)
    await db.commit()
    return depts


@pytest_asyncio.fixture
async def seed_persons(db, seed_departments):
    persons = [
        Person(id="person-1", name="张三", department_id="dept-tech", department_name="技术部", position="工程师"),
        Person(id="person-2", name="李四", department_id="dept-product", department_name="产品部", position="产品经理"),
        Person(id="person-3", name="王五", department_id="dept-tech", department_name="技术部", position="架构师"),
        Person(id="person-4", name="赵六", department_id=None, department_name="", position="实习生"),
    ]
    for p in persons:
        db.add(p)
    await db.commit()
    return persons


@pytest_asyncio.fixture
async def seed_scoring_config(db):
    config = ScoringConfig(
        id="default",
        name="默认评分配置",
        dimensions=[
            {"name": "工作闭环", "weight": 40, "max_score": 100, "description": "完成度"},
            {"name": "产出对比", "weight": 30, "max_score": 100, "description": "可量化"},
            {"name": "行文规范", "weight": 30, "max_score": 100, "description": "结构清晰"},
        ],
        grade_thresholds={"S": 90, "A": 80, "B": 70, "C": 60},
        prompt_template="",
        report_prompt="你是一位周报评分专家。请根据周报内容评分，满分100分。输出JSON格式。",
    )
    db.add(config)
    await db.commit()
    return config


@pytest_asyncio.fixture
async def seed_ai_model(db):
    """Seed 一个可用的 AI 模型配置（qwen3.7-plus）"""
    from app.models.models import AIModel
    model = AIModel(
        id="test-ai-model",
        name="qwen3.7-plus",
        provider="openai",
        model_id="qwen3.7-plus",
        api_key="sk-ws-H.RXLLPYD.EZ2r.MEQCIFAIRQ9LLE2dTV9n2FTjqAevbyNeCO1akYmm2bIF6CHJAiA1SdMtXf0BfYx5lexqqyR7_InCpdyQXT1nXmx1kkm9Nw",
        base_url="https://llm-5e0l0navgirl2i2v.cn-beijing.maas.aliyuncs.com/compatible-mode/v1",
        is_vision=True,
        is_active=True,
    )
    db.add(model)
    await db.commit()
    return model


def get_current_week():
    today = date.today()
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    return monday, sunday


def make_excel_file(
    filepath: str,
    last_week_title: str = None,
    this_week_title: str = None,
    last_week_rows: list = None,
    this_week_rows: list = None,
    empty: bool = False,
):
    """创建测试用 Excel 文件"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "测试周报"

    if empty:
        wb.save(filepath)
        return

    monday, sunday = get_current_week()
    last_monday = monday - timedelta(days=7)
    last_sunday = sunday - timedelta(days=7)

    if last_week_title is None:
        last_week_title = f"上周工作内容：{last_monday.strftime('%Y.%m.%d')}-{last_sunday.strftime('%Y.%m.%d')}"

    if this_week_title is None:
        this_week_title = f"本周工作计划：{monday.strftime('%Y.%m.%d')}-{sunday.strftime('%Y.%m.%d')}"

    headers = ["序号", "客户/项目", "工作内容", "汇报人", "结果反馈"]

    ws.merge_cells("A1:E1")
    ws["A1"] = last_week_title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = Font(bold=True)

    if last_week_rows is None:
        last_week_rows = [
            [1, "AI预警系统", "优化AI预警系统宣传资料", "张三", "完成优化"],
            [2, "激光雷达", "参与激光雷达测试", "张三", "发现识别问题"],
        ]

    for row_idx, row_data in enumerate(last_week_rows, 3):
        for col_idx, val in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    gap_row = 3 + len(last_week_rows)

    ws.merge_cells(f"A{gap_row}:E{gap_row}")
    ws[f"A{gap_row}"] = this_week_title
    ws[f"A{gap_row}"].font = Font(bold=True, size=14)
    ws[f"A{gap_row}"].alignment = Alignment(horizontal="center")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=gap_row + 1, column=col, value=h)
        cell.font = Font(bold=True)

    if this_week_rows is None:
        this_week_rows = [
            [1, "考勤系统", "完成评分系统demo", "张三", "完成demo"],
        ]

    for row_idx, row_data in enumerate(this_week_rows, gap_row + 2):
        for col_idx, val in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    wb.save(filepath)


def make_docx_file(filepath: str, content: str = None):
    """创建测试用 Word 文件"""
    from docx import Document

    doc = Document()
    if content:
        doc.add_paragraph(content)
    else:
        monday, sunday = get_current_week()
        last_monday = monday - timedelta(days=7)
        last_sunday = sunday - timedelta(days=7)
        doc.add_paragraph(f"上周工作内容：{last_monday.strftime('%Y.%m.%d')}-{last_sunday.strftime('%Y.%m.%d')}")
        doc.add_paragraph("1. 完成AI预警系统优化 2. 参与激光雷达测试")
        doc.add_paragraph(f"本周工作计划：{monday.strftime('%Y.%m.%d')}-{sunday.strftime('%Y.%m.%d')}")
        doc.add_paragraph("1. 完成考勤评分系统demo")
    doc.save(filepath)


def make_empty_file(filepath: str, ext: str):
    """创建空文件或无效文件"""
    if ext == ".xlsx":
        import openpyxl
        wb = openpyxl.Workbook()
        wb.save(filepath)
    elif ext == ".docx":
        from docx import Document
        Document().save(filepath)
    else:
        with open(filepath, "w") as f:
            f.write("")
